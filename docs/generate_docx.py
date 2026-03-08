#!/usr/bin/env python3
"""
Generates docs/documentacao.docx from docs/documentacao.md using python-docx.
Run from the repo root:  python3 docs/generate_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


MD = Path(__file__).parent / "documentacao.md"
DOCX = Path(__file__).parent / "documentacao.docx"


def set_cell_bg(cell, hex_color: str):
    """Set background fill colour of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int):
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str):
    """Add a paragraph, handling inline bold (**text**)."""
    p = doc.add_paragraph()
    _apply_inline(p, text)


def _apply_inline(paragraph, text: str):
    """Split on **bold** markers and add runs accordingly."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # handle inline code  `code`
            sub = re.split(r"(`[^`]+`)", part)
            for s in sub:
                if s.startswith("`") and s.endswith("`"):
                    run = paragraph.add_run(s[1:-1])
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                else:
                    paragraph.add_run(s)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()  # remove default empty run
    _apply_inline(p, text)


def add_code_block(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph(style="No Spacing")
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        # light grey background per paragraph
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)


def add_table(doc: Document, header: list[str], rows: list[list[str]]):
    col_count = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
        set_cell_bg(hdr_cells[i], "4B44CC")
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        fill = "FFFFFF" if r_idx % 2 == 0 else "F0EFFF"
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = cell_text
            set_cell_bg(row_cells[c_idx], fill)

    doc.add_paragraph()  # spacing after table


def parse_and_build(md_text: str):
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(1.18)

    lines = md_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Horizontal rule ───────────────────────────────────────────────────
        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        if line.startswith("#### "):
            add_heading(doc, line[5:], level=4)
            i += 1
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:], level=3)
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:], level=2)
            i += 1
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:], level=1)
            i += 1
            continue

        # ── Fenced code block ─────────────────────────────────────────────────
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            add_code_block(doc, code_lines)
            continue

        # ── Table ─────────────────────────────────────────────────────────────
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1

            # first line = header, second line = separator, rest = rows
            def split_row(r):
                return [c.strip() for c in r.strip().strip("|").split("|")]

            if len(table_lines) >= 2:
                header = split_row(table_lines[0])
                data_rows = [split_row(r) for r in table_lines[2:]]
                add_table(doc, header, data_rows)
            continue

        # ── Bullet list ───────────────────────────────────────────────────────
        if line.startswith("- "):
            add_bullet(doc, line[2:])
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────────────
        if line.strip() == "":
            i += 1
            continue

        # ── Regular paragraph ─────────────────────────────────────────────────
        add_paragraph(doc, line)
        i += 1

    return doc


def main():
    md_text = MD.read_text(encoding="utf-8")
    doc = parse_and_build(md_text)
    doc.save(str(DOCX))
    print(f"Saved: {DOCX}")


if __name__ == "__main__":
    main()
